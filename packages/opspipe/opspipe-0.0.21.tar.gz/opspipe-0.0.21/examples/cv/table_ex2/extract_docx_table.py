# -*- coding: utf-8 -*-
'''
Created on 2020-11-11

@author: zhys513(254851907@qq.com)
''' 
#pip install python-docx
from docx.api import Document
import pandas as pd

def extract_docx_table(input_file, output_file):
    """'
    docx的表格提取保存到excel
    :param input_file: 输入的docx文件
    :param output_file: 保存table到excel文件
    :return:
    """
    doc = Document(input_file)
    writer = pd.ExcelWriter(output_file, engine='xlsxwriter')
    # 如果有多个表格，对每个表格都保存成一个sheet
    for i in range(len(doc.tables)):
        table = doc.tables[i]
        data = []
        keys = None
        for j, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if j == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        df = pd.DataFrame(data)
        #保存到sheet
        df.to_excel(writer, sheet_name='table_{}'.format(i))
    writer.save()

if __name__ == '__main__': 
    
    filename = '石景山区高井规划一路道路工程广宁路15#杆输变电工程勘察设计合同'
    input_file = f"test/{filename}.docx"
    output_file = f"out/{filename}.xlsx"
    extract_docx_table(input_file, output_file)