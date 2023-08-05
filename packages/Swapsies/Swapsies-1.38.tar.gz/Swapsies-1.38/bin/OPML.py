#!/usr/bin/env python3

# PYTHON_ARGCOMPLETE_OK

import os,sys,re,json,codecs,xmltodict,unicodedata

from collections import OrderedDict
from io import StringIO
from zipfile import ZipFile

from Perdy.pyxbext import directory

from xlrd import open_workbook,xldate_as_tuple
from xlwt import Workbook

from docx import Document
from docx.styles.style import _ParagraphStyle

from Argumental.Argue import Argue
from Baubles.Logger import Logger
from Baubles.Colours import Colours
from Perdy.parser import printXML
from Perdy.pretty import prettyPrintLn, Style
from GoldenChild.xpath import *

args=Argue()
logger=Logger()
colours=Colours()

#_________________________________________________________________
@args.command(single=True)
class OPML(object):
	'''
	opml conversion tool
	'''

	#_____________________________________________________________
	def _clean(self, text):
		if not text:
			return ''
		if type(text) is str:
			text = unicode(text)
		return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore')


	#.............................................................
	def _load(self, file, extension=None):
		(self.doc,self.ctx,self.nsp) = getContextFromFile(file)
		if extension:
			backup = '%s.%s'%(file,extension)
			if os.path.isfile(backup):
				os.unlink(backup)
			os.rename(file, backup)


	#.............................................................
	def _save(self, file):
		with codecs.open(file, 'w', encoding='UTF-8') as output:
			printXML(str(self.doc), output = output)
			

	#.............................................................
	@args.operation
	def opml2org(self, file):
		'''
		from cloud outliner to emacs org mode
		'''
		self._load(file,'outline')
		xH = '//outline[@text]'
		for headline in getElements(self.ctx,xH):
			setAttribute(headline, 'structure', 'headline')
			note = getAttribute(headline, '_note')
			delAttribute(headline, '_note')
			if not note: continue
			for paragraph in note.split('&#10;'):
				child = addElement(self.doc,'outline',headline)
				setAttribute(child,'structure','paragraph')
				setAttribute(child,'text', paragraph)
		self._save(file)


	#.............................................................
	def _recurse(self, sheet, parent, row=0, col=0):
		for child in getElements(self.ctx, 'outline', parent):
			try:
				text = getAttribute(child, 'text')
				text = self._clean(text)
			except:
				text = '?'
			#print(row, col, text)
			sheet.write(row,col,text)
			try:
				note = getAttribute(child, '_note')
				node = self._clean(node)
			except:
				note = None
			if note:
				row += 1
				sheet.write(row,col+1,note)
			row = self._recurse(sheet, child, row+1, col+1)
		return row


	#.............................................................
	@args.operation
	def opml2excel(self, file):
		'''
		convert opml to excel format
		'''
		if not file.endswith('opml'):
			sys.stderr.write('not an opml file = %s\n'%file)
			return

		self._load(file)
		workbook = Workbook()
		sheet = workbook.add_sheet('opml')
		self._recurse(sheet, getElement(self.ctx, '/opml/body'))
		workbook.save(file.replace('.opml','.xls'))


	#.............................................................
	@args.operation
	def org2opml(self, file):
		'''
		from emacs org mode to cloud outliner
		'''
		self._load(file,'org')
		xH = '//outline[@structure="headline"]'
		xP = 'outline[@structure="paragraph"]'
		for headline in getElements(self.ctx,xH):
			delAttribute(headline, 'structure')
			sio = StringIO()
			for paragraph in getElements(self.ctx,xP,headline):
				sio.write('%s\n'%paragraph.prop('text'))
				paragraph.unlinkNode()
				paragraph.freeNode()
			note = sio.getvalue()
			setAttribute(headline, '_note', note)
		self._save(file)


	#.............................................................
	@args.operation
	def text2opml(self, file):
		'''
		read a text file using the indent as the level
		'''
		textPattern   = re.compile('^(\s+)(\S\S.*)$')
		bulletPattern = re.compile('^(\S[\.\)])(\s*)(\S.*)$')
		
		with open(file) as input:
			for line in input.read().split('\n'):
				line = line.rstrip('\n')
				textMatch = textPattern.match(line)
				if textMatch:
					(indent, text) = textMatch.groups()
				else:
					bulletMatch = bulletPattern.match(line)
					if bulletMatch:
						(indent, text) = bulletMatch.groups()[1:]
					else:
						continue
				print('%s%s'%(indent, text))


	#.............................................................
	@args.operation
	def json2opml(self, file):
		'''
		read a json file and convert to opml
		'''

		ext = '.%s'%(file.split('.')[-1].lower())
		if ext != '.json':
			sys.stderr.write('not a json file %s\n'%file)
			return
		
		def dig(node, parent):
			if type(node) in [list]:
				parent['outline'] = []
				for i in range(len(node)):
					item = node[i]
					child = OrderedDict([
						('@text','[%s]'%i)
					])
					dig(item, child)
					parent['outline'].append(child)
				return 
			
			if type(node) in [dict,OrderedDict]:
				parent['outline'] = []
				for name, value in node.items():
					child = OrderedDict([
						('@text', name),
					])
					dig(value, child)
					parent['outline'].append(child)
				return
			
			# assume fundamental fromhere
			parent['@_note'] = node


		opml = OrderedDict([
			('opml', OrderedDict([
				('@version', '1.0'),
				('head', OrderedDict([
					('title', file),
				])),
				('body', OrderedDict([
				])) 
			]))
		])
		
		with open(file) as input:
			d = json.load(input)
			dig(d, opml['opml']['body'])
			#print(json.dumps(opml,indent=4))
				  
			with open(file.replace(ext,'.opml'),'w') as output:
				xmltodict.unparse(opml, output)
					  

	#.............................................................
	def _createOPML(self, title):
		opml = OrderedDict([
			('opml', OrderedDict([
				('@version', '1.0'),
				('head', OrderedDict([
					('title', title.replace('&','and')),
					('expansionState', '0'),
				])),
				('body', OrderedDict([
				])) 
			]))
		])
		return opml
	
		
	#.............................................................
	@args.operation
	def docxComments2opml(self, file):
		if not file.endswith('docx'):
			sys.stderr.write('not a docx file\n')
			return

		path = os.path.expanduser(file)
		zip = ZipFile(path)
		xml = zip.read('word/comments.xml').decode('UTF8')
		doc = XML(*getContextFromString(xml))

		opml = self._createOPML(file)	
		body = opml['opml']['body']
		body['outline'] = []

		for comment in getElements(doc.ctx, '//w:comment'):
			print(comment.content)
			body['outline'].append({
				'@text': comment.content,
				'@_note': '%s %s'%(
					getAttribute(comment, 'date').replace('T',' ').replace('Z',''),
					getAttribute(comment, 'author')
				),
			})					

		name = re.sub('.docx*$', '.opml', file, flags=re.IGNORECASE)
		with open(name, 'w') as output:
			xmltodict.unparse(opml, output, encoding='UTF8', pretty=True)
		print(name)			

		
	#.............................................................
	@args.operation
	@args.parameter(name="note", short='n', flag=True, help='insert as notes not element')
	def docxHeadings2opml(self, file, note=False):
		if not file.endswith('docx'):
			sys.stderr.write('not a docx file\n')
			return

		path =os.path.expanduser(file)
		doc = Document(path)

		opml = self._createOPML(file)
		
		body = opml['opml']['body']
		body['outline'] = []
		stack = [ body ]
		level = 0
		
		highlites = { 
			1: colours.Green, 
			2: colours.Blue, 
			3: colours.White 
		}

		for paragraph in doc.paragraphs:
			ps = paragraph.style
			text = paragraph.text.strip().replace('\t',' ').replace('\n',' ')
			if len(text) == 0:
				continue

			parent = stack[-1]

			if ps.name.startswith('Heading'):
				heading = ps.name.replace('Heading ','').split(' ')[0]
				level=int(heading)

				if level in highlites.keys():
					sys.stdout.write(highlites[level])
				sys.stdout.write('%s%s\n'%('  '*(level-1), text))

				if level < len(stack):
					parent = stack.pop()
				if level > len(stack):
					parent['outline'] = []
					stack.append(parent)
					parent = stack[-1]
					
				parent['@text'] = text

			else:
				sys.stdout.write('%s%s\n'%('  '*(level), text))
				
				
			sys.stdout.write(colours.Off)

			if 'outline' not in parent.keys():
				parent['outline'] = []
			if '@_note' not in parent.keys():
				parent['@_note'] = ''

			if note:
				lines = list(filter(lambda x: len(x), parent['@_note'].split('\n')))
				lines.append(text)
				parent['@_note'] = '\n'.join(lines)
			else: 
				parent['outline'].append({
					'@text': text,
				})
					
		
		name = re.sub('.docx*$', '.opml', file, flags=re.IGNORECASE)
		with open(name, 'w') as output:
			xmltodict.unparse(opml, output, encoding='UTF8', pretty=True)
		print(name)


	#.............................................................
	@args.operation
	def docxIndented2opml(self, file):
		if not file.endswith('docx'):
			sys.stderr.write('not a docx file\n')
			return

		path = os.path.expanduser(file)
		doc = Document(path)

		opml = self._createOPML(file)
		
		body = opml['opml']['body']
		# stack[-1] is the current parent
		stack = [ body ]

		p = re.compile('^(\t*| *)(\S.*)$')
		
		level=1

		for paragraph in doc.paragraphs:
			ps = paragraph.style
			pf = ps.paragraph_format

			text = paragraph.text.strip()
			
			print(
				paragraph.text,
				paragraph.paragraph_format.left_indent,
				paragraph.style.paragraph_format.left_indent,
			)

			if len(text) == 0: continue

			outline = {
				'@text': text,
			}

			if level > len(stack):
				parent = stack[-1]['outline'][-1]
				stack.append(parent)

			if level < len(stack):
				stack.pop()
				
			if 'outline' not in stack[-1].keys():
				stack[-1]['outline'] = []
			stack[-1]['outline'].append(outline)				

			
		name = re.sub('.docx*$', '.opml', file, flags=re.IGNORECASE)
		with open(name,'w') as output:
			xmltodict.unparse(opml, output, encoding='UTF8', pretty=True)
		print(name)


	#.............................................................
	@args.operation
	def xlsx2opml(self, file):
		if not file.endswith('xlsx'):
			sys.stderr.write('not an xlsx file\n')
			return

		path = os.path.expanduser(file)
		workbook = open_workbook(filename=path)

		opml = self._createOPML(file)
		
		body = opml['opml']['body']

		# stack[-1] is the current parent
		stack = [ body ]

		def _push(stack, value):
			if 'outline' not in stack[-1].keys():
				stack[-1]['outline'] = []
			outline = {
				'@text': value,
			}
			stack[-1]['outline'].append(outline)				
			parent = stack[-1]['outline'][-1]
			stack.append(parent)
				
		def _pop(stack):
			stack.pop()

		for sheet in workbook.sheets():

			_push(stack, sheet.name)

			for r in range(sheet.nrows):
			
				for c in range(sheet.ncols):
					_push(stack, sheet.cell(r,c).value)

				for c in range(sheet.ncols):
					_pop(stack)
					
		name = re.sub('.xlsx*$', '.opml', file, flags=re.IGNORECASE)
		with open(name,'w') as output:
			xmltodict.unparse(opml, output, encoding='UTF8', pretty=True)
		print(name)


#_________________________________________________________________
if __name__ == '__main__': args.execute()

