#!/usr/bin/env python3

# PYTHON_ARGCOMPLETE_OK

import sys, os, json, logging, arrow, codecs

#sys.path.insert(0,'..')

from datetime import datetime
from uuid import uuid4

from docx import Document
from docx.shared import Inches, RGBColor

from Baubles.Logger import Logger
from Perdy.pretty import prettyPrint
from Perdy.parser import printXML
from Perdy.pyxbext import directory
from Argumental.Argue import Argue
from Swapsies.xmi import XMI
from GoldenChild.xpath import *

logger = Logger()
args = Argue()

logger.setLevel(logging.INFO)

#____________________________________________________________
@args.command(single=True)
class COD(object):

	#........................................................
	def __init__(self):
		'''
		setup constructor with lookup lists
		'''

		self.xcolours = {
			'White': '\033[37m',
			'Red': '\033[31m',
			'Orange': '\033[33m',
			'Green': '\033[32m',
			'Teal': '\033[36m',
			'Blue': '\033[34m',
			'Purple': '\033[35m',
			'Black': '\033[30m',
			'Off': '\033[0m',
		}

		self.colours = {
			0: 'White',
			#1: 'Grey', 
			2: 'Red',
			3: 'Orange',
			4: 'Green',
			5: 'Blue',
			6: 'Purple',
		}

		self.xfonts = {
			'normal': '\033[0m',
			'strikeout': '\033[9m',
			'underline': '\033[4m',
			'italics': '\033[3m',
			'bold': '\033[1m',
		}

		self.fonts = {
			0: 'normal',
			1: 'strikeout',
			2: 'underline',
			4: 'italics',
			8: 'bold',
		}

	
	#........................................................
	def __attributes(self, cod, childItem, checkboxes):
		font = int(getElementText(cod.ctx, 'fontStyle', childItem))
		colour = int(getElementText(cod.ctx, 'color', childItem))
		state = int(getElementText(cod.ctx, 'completionState', childItem)) == 3

		if checkboxes:
			checked = '✓ ' if state else '□ '
		else:
			checked = ''

		return font, colour, state, checked
	

	#........................................................
	def __document(self, cod, file):
		s = getElementText(cod.ctx, '/Document/Properties/lastModificationTime')
		u = arrow.get(float(s))
		a = u.to('local').format('YYYY-MM-DD HH:mm:ss SSS Z')
		t = getElementText(cod.ctx, '/Document/Properties/title')
		print('%s -> "%s" => %s' % (file, t, a))

				   
	#........................................................
	@logger.debug
	def __node2text(self, cod, childItem, indent='', checkboxes=False, shownotes=False):
		'''
		show a single item
		'''

		font, colour, state, checked = self.__attributes(cod, childItem, checkboxes)
		
		if colour in self.colours.keys():
			_colour = self.xcolours[self.colours[colour]] or ''
		else:
			_colour = ''
		_font = self.xfonts[self.fonts[font]] or ''
		
		print('%s%s%s%s%s%s' % (
			indent, _font, _colour, checked,
			getElementText(cod.ctx, 'title', childItem),
			self.xcolours['Off'], )
		)

		if shownotes:
			notes = getElement(cod.ctx, 'notes', childItem)
			if notes:
				note = notes.content
				if note != '(null)':
					print('%s%s  "%s"%s' % (
						self.xcolours['Teal'], '%s  ' % indent
						if checkboxes else indent, note,
						self.xcolours['Off'],
					))

		for grandChild in getElements(cod.ctx, 'ChildItem', childItem):
			self.__node2text(
				cod,
				grandChild,
				indent='%s  ' % indent,
				checkboxes=checkboxes,
				shownotes=shownotes
			)
			

	#........................................................
	@logger.debug
	def __nodes2text(self, file, cod, checkboxes=False, shownotes=False):
		'''
		show the cod file
		'''
		self.__document(cod, file)
		
		for childItem in getElements(cod.ctx, '/Document/Properties/context/ChildItem'):
			self.__node2text(cod, childItem, checkboxes=checkboxes, shownotes=shownotes)
			

	#........................................................
	@logger.debug
	@args.operation
	@args.parameter(name='checkboxes', short='c', flag=True, help='show checkboxes')
	@args.parameter(name='shownotes', short='n', flag=True, help='show text notes')
	@args.parameter(name='blackAndWhite', short='b', flag=True, help='black and white')
	def cod2text(self, file, checkboxes=False, shownotes=False, blackAndWhite=False):
		'''
		load a cod file and display
		'''
		#self.reorganize(file)
		
		if blackAndWhite:
			for key in list(self.xcolours.keys()):
				self.xcolours[key] = ''
			for key in list(self.xfonts.keys()):
				self.xfonts[key] = ''

		cod = XML(*getContextFromFile(os.path.expanduser(file)))
		self.__nodes2text(file, cod, checkboxes=checkboxes, shownotes=shownotes)


	#........................................................
	@logger.debug
	def __node2uml(self, cod, xmi, parent, diagram, childItem, indent='', checkboxes=False, shownotes=False):
		'''
		show a single item
		'''

		font, colour, state, checked = self.__attributes(cod, childItem, checkboxes)
		
		if colour in self.colours.keys():
			_colour = self.xcolours[self.colours[colour]] or ''
		else:
			_colour = ''
		_font = self.xfonts[self.fonts[font]] or ''

		text = getElementText(cod.ctx, 'title', childItem)
		
		print('%s%s%s%s%s%s' % (
			indent, _font, _colour, checked,
			text,
			self.xcolours['Off'], )
		)

		note = None
		if shownotes:
			notes = getElement(cod.ctx, 'notes', childItem)
			if notes:
				note = notes.content
				if note != '(null)':
					print('%s%s  "%s"%s' % (
						self.xcolours['Teal'], '%s  ' % indent
						if checkboxes else indent, note,
						self.xcolours['Off'],
					))
				else:
					note  = None

		children = getElements(cod.ctx, 'ChildItem', childItem)
		
		if len(children) == 0:
			_child = xmi.makeClass(text, parent)
		else:
			_child = xmi.makePackage(text, parent)
			_diagram = xmi.makeClassDiagram(text, _child)

			child = None
			for grandChild in children:
				sibling = self.__node2uml(
					cod,
					xmi,
					_child,
					_diagram,
					grandChild,
					indent='%s  ' % indent,
					checkboxes=checkboxes,
					shownotes=shownotes
				)

				if child:
					xmi.makeDependency(child, sibling, parent)
				child = sibling

		xmi.addDiagramClass(_child, diagram)

		if note:
			xmi.makeNote(note, _child)

		xmi.makeContentTag('completed', checked, _child)
		
		return _child
	

	#........................................................
	@logger.debug
	def __nodes2uml(self, file, cod, xmi, checkboxes=False, shownotes=False):
		'''
		show the cod file
		'''
		self.__document(cod, file)
		
		diagram = xmi.makeClassDiagram(file, xmi.modelNS)
		for childItem in getElements(cod.ctx, '/Document/Properties/context/ChildItem'):
			self.__node2uml(cod, xmi, xmi.modelNS, diagram, childItem, checkboxes=checkboxes, shownotes=shownotes)
			

	#........................................................
	@logger.debug
	@args.operation
	@args.parameter(name='checkboxes', short='c', flag=True, help='show checkboxes')
	@args.parameter(name='shownotes', short='n', flag=True, help='show text notes')
	def cod2uml(self, file, checkboxes=False, shownotes=False):
		'''
		convert a COD file to Sparx UML packages
		'''

		cod = XML(*getContextFromFile(os.path.expanduser(file)))
		xmi = XMI(file)
		
		self.__nodes2uml(file, cod, xmi, checkboxes=checkboxes, shownotes=shownotes)

		with open(f'{file}.xmi', 'w') as output:
			output.write(str(xmi.doc))


	#........................................................
	@logger.debug
	def __node2docx(self, cod, docx, childItem, depth=None, checkboxes=False, shownotes=False, blackAndWhite=False):
		'''
		show a single item
		'''

		font, colour, state, checked = self.__attributes(cod, childItem, checkboxes)

		if colour in self.colours.keys():
			_colour = self.xcolours[self.colours[colour]] or ''
		else:
			_colour = ''
		_font = self.xfonts[self.fonts[font]] or ''
		
		if not depth: depth = 1
		
		text = getElementText(cod.ctx, 'title', childItem)
		print('%s%s%s%s%s%s' % (
			'  '*(depth-1), _font, _colour, checked,
			text,
			self.xcolours['Off'], )
		)

		run = docx.add_heading(level=depth).add_run(checked+text)
		if font & 8: run.bold = True
		if font & 4: run.italic = True
		if font & 2: run.underline = True
		if font & 1: run.font.strike = True

		if not blackAndWhite:
			if colour == 1: run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
			if colour == 2: run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
			if colour == 3: run.font.color.rgb = RGBColor(0xff, 0x80, 0x00)
			if colour == 4: run.font.color.rgb = RGBColor(0x00, 0xff, 0x00)
			if colour == 5: run.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
			if colour == 6: run.font.color.rgb = RGBColor(0xff, 0x00, 0xff)

		if shownotes:
			notes = getElement(cod.ctx, 'notes', childItem)
			if notes:
				note = notes.content
				if note != '(null)':
					print('%s%s  "%s"%s' % (
						self.xcolours['Teal'], '%s  ' % ('  '*(depth-1))
						if checkboxes else '  '*(depth-1), note,
						self.xcolours['Off'],
					))
					indent = Inches((depth-1)*0.25)
					docx.add_paragraph(note).paragraph_format.left_indent = indent


		for grandChild in getElements(cod.ctx, 'ChildItem', childItem):
			self.__node2docx(
				cod,
				docx,
				grandChild,
				depth=depth+1,
				checkboxes=checkboxes,
				shownotes=shownotes,
				blackAndWhite=blackAndWhite,
			)
			

	#........................................................
	@logger.debug
	def __nodes2docx(self, file, cod, docx, checkboxes=False, shownotes=False, blackAndWhite=False):
		'''
		show the cod file
		'''
		self.__document(cod, file)
		for childItem in getElements(cod.ctx, '/Document/Properties/context/ChildItem'):
			self.__node2docx(cod, docx, childItem, checkboxes=checkboxes, shownotes=shownotes, blackAndWhite=blackAndWhite)
			

	#........................................................
	@logger.debug
	@args.operation
	@args.parameter(name='checkboxes', short='c', flag=True, help='show checkboxes')
	@args.parameter(name='shownotes', short='n', flag=True, help='show text notes')
	@args.parameter(name='blackAndWhite', short='b', flag=True, help='black and white')
	def cod2docx(self, file, checkboxes=False, shownotes=False, blackAndWhite=False):
		'''
		load a cod file and display
		'''
		#self.reorganize(file)

		docx = Document()

		p = re.compile('^Heading (\d)$')

		for style in docx.styles:
			m =  p.match(style.name)
			if m:
				depth = int(m.group(1))
				indent = Inches((depth-1)*0.25)
				style.paragraph_format.left_indent = indent

		if blackAndWhite:
			for key in list(self.xcolours.keys()):
				self.xcolours[key] = ''
			for key in list(self.xfonts.keys()):
				self.xfonts[key] = ''

		cod = XML(*getContextFromFile(os.path.expanduser(file)))
		self.__nodes2docx(file, cod, docx, checkboxes=checkboxes, shownotes=shownotes, blackAndWhite=blackAndWhite)

		docx.save(f'{file}.docx')


#____________________________________________________________
if __name__ == '__main__': args.execute()


